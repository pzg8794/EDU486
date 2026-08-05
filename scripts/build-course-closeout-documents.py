#!/usr/bin/env python3
"""Build accessible DOCX closeout packets from reviewed Markdown sources."""

from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TMP = ROOT / "tmp" / "course-closeout-documents"
SUBMISSIONS = ROOT / "public-submissions"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def create_reference(path: Path, landscape: bool) -> None:
    default_reference = subprocess.run(
        ["/opt/homebrew/bin/pandoc", "--print-default-data-file", "reference.docx"],
        check=True,
        stdout=subprocess.PIPE,
    ).stdout
    path.write_bytes(default_reference)
    doc = Document(path)
    section = doc.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.68)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.90)
    section.footer_distance = Inches(0.25)

    styles = doc.styles
    for name in ("Normal", "Body Text", "First Paragraph", "Compact"):
        if name not in styles:
            continue
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.08

    colors = {
        "Title": "123047",
        "Heading 1": "123047",
        "Heading 2": "087F8C",
        "Heading 3": "66338C",
    }
    sizes = {"Title": 24, "Heading 1": 20, "Heading 2": 15, "Heading 3": 12}
    for name, color in colors.items():
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(sizes[name])
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    doc.save(path)


def prepare_part2_markdown(source: Path, destination: Path) -> None:
    text = source.read_text(encoding="utf-8")
    image = (ROOT / "public-artifacts" / "grounding-frameworks-part2-visual.png").as_posix()
    replacement = (
        "![Concentric JuST model with a justice-centered community, a "
        "Plan-Elicit-Revise-Act cycle, a youth-evidence redesign loop, and "
        f"technology support, complication, and constraint lenses.]({image})"
    )
    text = re.sub(
        r"\[!\[[^\]]*\]\(grounding-frameworks-part2-visual\.svg\)\]"
        r"\(grounding-frameworks-part2-visual\.svg\)",
        replacement,
        text,
        count=1,
    )
    destination.write_text(text, encoding="utf-8")


def run_pandoc(source: Path, output: Path, reference: Path) -> None:
    subprocess.run(
        [
            "/opt/homebrew/bin/pandoc",
            str(source),
            "--from=gfm",
            "--to=docx",
            f"--reference-doc={reference}",
            "--resource-path",
            str(ROOT),
            "--output",
            str(output),
        ],
        cwd=ROOT,
        check=True,
    )


def polish_docx(path: Path, title: str, alt_text: str | None = None) -> None:
    doc = Document(path)
    doc.core_properties.title = title
    doc.core_properties.subject = "EDU486 course closeout submission"
    doc.core_properties.author = "Piter Z. Garcia Bautista"
    doc.core_properties.keywords = "EDU486, justice-centered STEM, accessibility"

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.widow_control = True
        if paragraph.style.name == "Title":
            paragraph.paragraph_format.keep_with_next = True
        for run in paragraph.runs:
            if run.font.name is None:
                run.font.name = "Arial"

    for table in doc.tables:
        if "Table Grid" in doc.styles:
            table.style = "Table Grid"
        table.autofit = True
        if table.rows:
            for cell in table.rows[0].cells:
                set_cell_shading(cell, "DDEBF1")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string("123047")
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    if alt_text:
        for shape in doc.inline_shapes:
            shape.width = Inches(9.2)
            shape.height = Inches(6.325)
        for drawing in doc.element.body.iter(qn("wp:docPr")):
            drawing.set("title", "Grounding Frameworks Part 2 visual")
            drawing.set("descr", alt_text)

    doc.save(path)


def build() -> None:
    TMP.mkdir(parents=True, exist_ok=True)
    SUBMISSIONS.mkdir(parents=True, exist_ok=True)

    part2_source = ROOT / "assignments" / "01-grounding-frameworks" / "grounding-frameworks-part2.md"
    addendum_source = ROOT / "assignments" / "03-camp-unit-plan" / "post-camp-implementation-revision-addendum.md"
    commitments_source = ROOT / "assignments" / "07-final-reflection" / "commitments-to-k12.md"
    visual = ROOT / "public-artifacts" / "grounding-frameworks-part2-visual.png"
    missing = [
        path
        for path in (part2_source, addendum_source, commitments_source, visual)
        if not path.exists()
    ]
    if missing:
        raise SystemExit("Missing required source: " + ", ".join(str(path) for path in missing))

    landscape_ref = TMP / "reference-landscape.docx"
    portrait_ref = TMP / "reference-portrait.docx"
    create_reference(landscape_ref, landscape=True)
    create_reference(portrait_ref, landscape=False)

    staged_part2 = TMP / "grounding-frameworks-part2.md"
    prepare_part2_markdown(part2_source, staged_part2)

    part2_output = SUBMISSIONS / "grounding-frameworks-part2.docx"
    addendum_output = SUBMISSIONS / "invisible-invaders-post-camp-revision-addendum.docx"
    commitments_output = SUBMISSIONS / "final-reflection-commitments-to-k12.docx"
    run_pandoc(staged_part2, part2_output, landscape_ref)
    run_pandoc(addendum_source, addendum_output, portrait_ref)
    run_pandoc(commitments_source, commitments_output, portrait_ref)

    polish_docx(
        part2_output,
        "Visualizing Our Grounding Frameworks - Part 2",
        (
            "A justice-centered community anchors a Plan-Elicit-Revise-Act cycle. "
            "Youth feedback returns to planning. A facilitator redesign loop keeps "
            "direct observation, youth statements, facilitator interpretation, and "
            "uncertainty distinct. Technology can support, complicate, or constrain."
        ),
    )
    polish_docx(
        addendum_output,
        "Invisible Invaders: Post-Camp Implementation Revision Addendum",
    )
    polish_docx(
        commitments_output,
        "Final Reflection: Commitments to K-12",
    )

    print(part2_output)
    print(addendum_output)
    print(commitments_output)


if __name__ == "__main__":
    build()
